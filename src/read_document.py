from docx import Document


def extract_document_text(file_path):
    doc = Document(file_path)

    text_parts = []

    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    return "\n".join(text_parts)


input_file = "input/Red Herring Prospectus.docx"

text = extract_document_text(input_file)

print(text)