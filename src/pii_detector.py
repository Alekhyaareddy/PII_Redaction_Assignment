import re
import spacy
from docx import Document


# ============================================================
# LOAD NER MODEL
# ============================================================

nlp = spacy.load("en_core_web_sm")


# ============================================================
# REGEX PATTERNS
# ============================================================

EMAIL_PATTERN = r"""
\b
[A-Za-z0-9._%+-]+
@
[A-Za-z0-9.-]+
\.[A-Za-z]{2,}
\b
"""


PHONE_PATTERN = r"""
(?<!\d)
(?:
    # Indian mobile number with +91 / 91
    (?:\+?91[\s-]?)
    [6-9]\d{9}

    |

    # Indian number with STD code
    (?:\+?91[\s-]?)
    \d{2}[\s-]?\d{4}[\s-]?\d{4}

    |

    # Indian number with STD code
    (?:\+?91[\s-]?)
    \d{2}[\s-]?\d{8}

    |

    # Indian mobile number without country code
    [6-9]\d{9}

    |

    # Indian landline e.g. 022-68052182
    0\d{2,4}[\s-]\d{6,8}
)
(?!\d)
"""


IP_PATTERN = r"""
(?<![\d.])
(?:
    (?:25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)
    \.
    (?:25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)
    \.
    (?:25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)
    \.
    (?:25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)
)
(?![\d.])
"""


CREDIT_CARD_PATTERN = r"""
(?<!\d)
(?:\d[ -]?){13,19}
(?!\d)
"""


SSN_PATTERN = r"""
(?<!\d)
(?!000|666|9\d\d)
\d{3}
[- ]
\d{2}
[- ]
\d{4}
(?!\d)
"""


# IMPORTANT:
# No (?ix) inside the pattern.
# We pass re.IGNORECASE | re.VERBOSE to finditer/findall instead.

DOB_PATTERN = r"""
(?:
    date\s+of\s+birth
    |
    dob
    |
    d\.o\.b\.
    |
    birth\s+date
)
\s*[:=-]?\s*
(
    \d{1,2}[/-]\d{1,2}[/-]\d{2,4}
    |
    \d{1,2}\s+
    (?:January|February|March|April|May|June|July|August|
       September|October|November|December)
    \s+\d{4}
    |
    (?:January|February|March|April|May|June|July|August|
       September|October|November|December)
    \s+\d{1,2},?\s+\d{4}
)
"""
ADDRESS_PATTERN = r"""
(?ix)
\b
\d{1,5}
(?:[-/]\d{1,5})?
\s+
[A-Za-z0-9.,'&()/-]+
(?:\s+[A-Za-z0-9.,'&()/-]+){0,8}
\s+
(?:Road|Rd|Street|St|Lane|Ln|Marg|Nagar|Colony|Complex|Industrial|Estate|
Park|Building|Tower|Floor|Block|Sector|Taluka|Village|Society|Layout)
\b
"""


# ============================================================
# DOCUMENT EXTRACTION
# ============================================================

def extract_document_text(file_path):
    """Extract text from paragraphs and tables in a DOCX file."""

    doc = Document(file_path)

    text_parts = []

    # Paragraphs
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if text:
            text_parts.append(text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:

                text = cell.text.strip()

                if text:
                    text_parts.append(text)

    return "\n".join(text_parts)


# ============================================================
# EMAIL DETECTION
# ============================================================

def detect_emails(text):

    matches = re.findall(
        EMAIL_PATTERN,
        text,
        re.VERBOSE
    )

    return list(dict.fromkeys(matches))


# ============================================================
# PHONE DETECTION
# ============================================================

def detect_phone_numbers(text):

    matches = re.findall(
        PHONE_PATTERN,
        text,
        re.VERBOSE
    )

    return list(dict.fromkeys(matches))


# ============================================================
# IP DETECTION
# ============================================================

def detect_ip_addresses(text):

    matches = re.findall(
        IP_PATTERN,
        text,
        re.VERBOSE
    )

    return list(dict.fromkeys(matches))


# ============================================================
# CREDIT CARD DETECTION
# ============================================================

def luhn_check(number):
    """
    Validate a possible credit card number using Luhn algorithm.
    """

    digits = [int(d) for d in number if d.isdigit()]

    if not 13 <= len(digits) <= 19:
        return False

    total = 0

    parity = len(digits) % 2

    for i, digit in enumerate(digits):

        if i % 2 == parity:

            digit *= 2

            if digit > 9:
                digit -= 9

        total += digit

    return total % 10 == 0


def detect_credit_cards(text):

    candidates = re.findall(
        CREDIT_CARD_PATTERN,
        text,
        re.VERBOSE
    )

    results = []

    for candidate in candidates:

        if luhn_check(candidate):
            results.append(candidate)

    return list(dict.fromkeys(results))


# ============================================================
# SSN DETECTION
# ============================================================

def detect_ssns(text):

    matches = re.findall(
        SSN_PATTERN,
        text,
        re.VERBOSE
    )

    return list(dict.fromkeys(matches))


# ============================================================
# DOB DETECTION
# ============================================================

def detect_dobs(text):

    matches = re.findall(
        DOB_PATTERN,
        text,
        re.VERBOSE | re.IGNORECASE
    )

    return list(dict.fromkeys(matches))


# ============================================================
# NAME DETECTION
# ============================================================

def detect_names(text):

    doc = nlp(text)

    names = []

    # Words that frequently cause false positives
    excluded_terms = {
        "reference rate",
        "selling shareholder",
        "bidder",
        "bidders",
        "bid amount",
        "offer price",
        "floor price",
        "cap price",
        "telephone",
        "website",
        "company",
        "registered broker",
        "share transfer",
        "share transfer agents",
        "key managerial",
        "key managerial personnel",
        "particulars",
        "date",
        "bill",
        "operational",
        "dfi",
        "i-sec",
        "deccan gymkhana",
        "buena monte",
        "telephone",
        "registered",
        "website",
        "company",
        "bank",
        "branch",
        "village",
        "schedule",
        "acknowledgement slip",
        "escrow collection bank",
        "individual bidders",
        "qib bidders",
        "upi circulars",
        "wilful defaulter",
        "air conditioning",
        "high voltage direct",
        "mega volt-amperes",
        "photo voltaic",
        "non-gaap measures",
    }

    # Location/business words
    location_terms = {
        "mumbai",
        "pune",
        "bandra",
        "east",
        "nagar",
        "marg",
        "road",
        "lane",
        "village",
        "taluka",
        "complex",
        "industrial",
        "bhavan",
        "hospital",
        "chambers",
        "showroom",
        "society",
        "house",
        "building",
        "floor",
        "office",
        "facility",
        "colony",
        "area",
    }

    for ent in doc.ents:

        if ent.label_ != "PERSON":
            continue

        name = ent.text.strip()

        name_lower = name.lower()

        # Must contain at least two words
        if len(name.split()) < 2:
            continue

        # Exact exclusions
        if name_lower in excluded_terms:
            continue

        # Location/business false positives
        words = set(name_lower.split())

        if words.intersection(location_terms):
            continue

        # Numbers should not appear in names
        if any(char.isdigit() for char in name):
            continue

        # Emails
        if "@" in name:
            continue

        # URLs
        if "http" in name_lower:
            continue

        if "www." in name_lower:
            continue

        # Slash-separated names are often mixed entities
        if "/" in name:
            continue

        # Company/entity suffixes
        company_terms = [
            "private limited",
            "limited",
            "company",
            "website",
            "huf",
            "llp",
            "bank",
            "trust",
            "foundation",
            "society",
        ]

        if any(term in name_lower for term in company_terms):
            continue

        names.append(name)

    # Remove duplicates while preserving order
    names = list(dict.fromkeys(names))

    return names


# ============================================================
# COMPANY DETECTION
# ============================================================

def detect_companies(text):

    doc = nlp(text)

    companies = []

    for ent in doc.ents:

        if ent.label_ in {"ORG"}:

            company = ent.text.strip()

            if len(company) < 3:
                continue

            companies.append(company)

    # Additional company-name patterns
    company_pattern = r"""
    \b
    [A-Z][A-Za-z&.,' -]{2,}
    \s+
    (?:
        Limited
        |
        Private Limited
        |
        Pvt\.?\s+Ltd\.?
        |
        LLP
        |
        Corporation
        |
        Bank
        |
        Industries
        |
        Foundation
        |
        Trust
    )
    \b
    """

    regex_companies = re.findall(
        company_pattern,
        text,
        re.VERBOSE
    )

    companies.extend(regex_companies)

    # Remove duplicates
    companies = list(dict.fromkeys(companies))

    return companies


# ============================================================
# PRINT RESULTS
# ============================================================

def print_results(title, items):

    print()
    print("=" * 60)
    print(title)
    print("=" * 60)

    print(f"Total detected: {len(items)}")

    if not items:

        print("None detected.")

        return

    for item in items:

        print(item)
def detect_addresses(text):
    matches = re.findall(ADDRESS_PATTERN, text, re.VERBOSE)
    return list(dict.fromkeys(match.strip() for match in matches))

# ============================================================
# MAIN
# ============================================================

if __name__ == "__main__":

    input_file = "input/Red Herring Prospectus.docx"

    print("Reading document...")

    text = extract_document_text(input_file)

    print()
    print("=" * 60)
    print("PII DETECTION RESULTS")
    print("=" * 60)

    # Detect everything
    emails = detect_emails(text)

    phones = detect_phone_numbers(text)

    ips = detect_ip_addresses(text)

    credit_cards = detect_credit_cards(text)

    ssns = detect_ssns(text)

    dobs = detect_dobs(text)

    names = detect_names(text)

    companies = detect_companies(text)
    addresses = detect_addresses(text)

print()
print("=" * 60)
print("PHYSICAL / MAILING ADDRESSES")
print("=" * 60)
print(f"Total detected: {len(addresses)}")

for address in addresses:
    print(address)

    # Print results
    print_results("EMAILS", emails)

    print_results("PHONE NUMBERS", phones)

    print_results("IP ADDRESSES", ips)

    print_results("CREDIT CARD NUMBERS", credit_cards)

    print_results("SSNs", ssns)

    print_results("DATES OF BIRTH", dobs)

    print_results("PERSON NAMES", names)

    print_results("COMPANIES", companies)

    # Summary
    print()
    print("=" * 60)
    print("DETECTION SUMMARY")
    print("=" * 60)

    print(f"Emails:          {len(emails)}")
    print(f"Phone numbers:   {len(phones)}")
    print(f"IP addresses:    {len(ips)}")
    print(f"Credit cards:    {len(credit_cards)}")
    print(f"SSNs:            {len(ssns)}")
    print(f"DOBs:            {len(dobs)}")
    print(f"Names:           {len(names)}")
    print(f"Companies:       {len(companies)}")

    print()
    print("Detection completed successfully.")