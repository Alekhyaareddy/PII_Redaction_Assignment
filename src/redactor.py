import re
from pathlib import Path

import spacy
from docx import Document


# ============================================================
# PATHS
# ============================================================

BASE_DIR = Path(__file__).resolve().parent.parent

INPUT_FILE = BASE_DIR / "input" / "Red Herring Prospectus.docx"
OUTPUT_FILE = BASE_DIR / "output" / "redacted_prospectus.docx"


# ============================================================
# LOAD NLP MODEL
# ============================================================

MODEL_NAME = "en_core_web_sm"

try:
    nlp = spacy.load(MODEL_NAME)
except OSError:
    nlp = None


def ensure_nlp_ready():
    if nlp is None:
        raise RuntimeError(
            "The spaCy model 'en_core_web_sm' is missing. "
            "Install it with: python -m spacy download en_core_web_sm"
        )


# ============================================================
# PII PATTERNS
# ============================================================

EMAIL_PATTERN = re.compile(
    r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"
)


PHONE_PATTERN = re.compile(
    r"""
    (?<!\d)
    (?:
        # +91 9876543210
        \+?\s*91[\s-]?[6-9]\d{9}

        |

        # +91 81081 14949
        \+?\s*91[\s-]?[6-9]\d{4}[\s-]\d{5}

        |

        # +91 (20) 6729 5100
        \+?\s*91[\s-]*\(\d{2,4}\)[\s-]*\d{4}[\s-]*\d{4}

        |

        # +91 20 45053237
        \+?\s*91[\s-]?\(?\d{2,4}\)?[\s-]?\d{4}[\s-]?\d{4}

        |

        # 9876543210
        [6-9]\d{9}

        |

        # 81081 14949
        [6-9]\d{4}[\s-]\d{5}

        |

        # 020-45053237
        0\d{2,4}[\s-]\d{6,8}
    )
    (?!\d)
    """,
    re.VERBOSE,
)


IP_PATTERN = re.compile(
    r"""
    (?<![\d.])
    (?:
        (?:25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)
        \.
        (?:25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)
        \.
        (?:25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)
        \.
        (?:25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)
    )
    (?![\d.])
    """,
    re.VERBOSE,
)


SSN_PATTERN = re.compile(
    r"(?<!\d)(?!000|666|9\d\d)\d{3}[- ]\d{2}[- ]\d{4}(?!\d)"
)


CREDIT_CARD_PATTERN = re.compile(
    r"(?<!\d)(?:\d[ -]?){13,19}(?!\d)"
)


DOB_PATTERN = re.compile(
    r"""
    (?:
        date\s+of\s+birth
        |
        d\.?o\.?b\.?
        |
        birth\s+date
    )
    \s*[:=-]?\s*
    (
        \d{1,2}[/-]\d{1,2}[/-]\d{2,4}

        |

        \d{1,2}\s+
        (?:January|February|March|April|May|June|July|August|
        September|October|November|December)
        \s+\d{4}

        |

        (?:January|February|March|April|May|June|July|August|
        September|October|November|December)
        \s+\d{1,2},?\s+\d{4}
    )
    """,
    re.VERBOSE | re.IGNORECASE,
)


# ============================================================
# REDACTION PLACEHOLDER
# ============================================================

REDACTION_TOKEN_PATTERN = re.compile(
    r"\[REDACTED_[A-Z]+_\d+\]"
)


# ============================================================
# ADDRESS DETECTION
# ============================================================

ADDRESS_WORDS = {
    "village",
    "taluka",
    "district",
    "plot",
    "house",
    "flat",
    "road",
    "street",
    "lane",
    "marg",
    "nagar",
    "building",
    "floor",
    "park",
    "complex",
    "estate",
    "society",
    "tower",
    "bunglow",
    "bungalow",
    "colony",
    "residency",
    "apartment",
    "housing",
    "industrial",
    "farms",
    "garden",
    "sector",
    "block",
    "wing",
    "avenue",
    "business centre",
    "business center",
    "near",
    "opposite",
    "off",
}


ADDRESS_CITIES = {
    "pune",
    "mumbai",
    "delhi",
    "bangalore",
    "bengaluru",
    "hyderabad",
    "chennai",
    "bhopal",
    "kolkata",
    "jaipur",
    "ahmedabad",
    "lucknow",
    "kochi",
    "visakhapatnam",
}


INDIAN_STATES = {
    "maharashtra",
    "karnataka",
    "telangana",
    "tamil nadu",
    "delhi",
    "madhya pradesh",
    "gujarat",
    "west bengal",
    "rajasthan",
    "uttar pradesh",
    "kerala",
    "andhra pradesh",
}


PIN_PATTERN = re.compile(
    r"\b\d{3}\s?\d{3}\b"
)


# A conservative fallback for short, number-led addresses that do not include
# a PIN code. It requires both an address word and a supported city, which
# avoids treating ordinary location references as mailing addresses.
ADDRESS_WITHOUT_PIN_PATTERN = re.compile(
    r"""
    (?:(?<=^)|(?<=[\n;:]))
    (?=[^\n]{0,100}\b(?:Village|Road|Street|Lane|Marg|Nagar|Building|\
        Complex|Estate|Society|Tower|Colony|Apartment|Housing|Industrial|\
        Sector|Block|Avenue)\b)
    \d{1,5}(?:[-/]\d{1,5})?
    (?:\s*,?\s*[A-Za-z][A-Za-z.'-]*){1,10}
    \s*,?\s*
    (?:Pune|Mumbai|Delhi|Bangalore|Bengaluru|Hyderabad|Chennai|Bhopal|\
       Kolkata|Jaipur|Ahmedabad|Lucknow|Kochi|Visakhapatnam)
    \b
    """,
    re.VERBOSE | re.IGNORECASE,
)


# ============================================================
# KNOWN COMPANIES
# ============================================================

KNOWN_COMPANIES = [
    "KSH INTERNATIONAL LIMITED",
    "Nuvama Wealth Management Limited",
    "ICICI Securities Limited",
    "MUFG Intime India Private Limited",
    "Link Intime India Private Limited",
    "Kirtane & Pandit LLP",
    "WATERLOO INDUSTRIAL PARK VI PRIVATE LIMITED",
]


# ============================================================
# KNOWN PEOPLE
# ============================================================

KNOWN_NAMES = [
    "Sarthak Malvadkar",
    "Kushal Subbayya Hegde",
    "Pushpa Kushal Hegde",
    "Rajesh Kushal Hegde",
    "Rohit Kushal Hegde",
    "Rakhi Girija Shetty",
    "Lokesh Shah",
    "Soumavo Sarkar",
    "Kishan Rastogi",
    "Abhijit Diwan",
    "Shanti Gopalkrishnan",
    "Sandesh Bhagwat",
    "Amod Joshi",
]


# ============================================================
# COUNTERS
# ============================================================

replacement_counters = {
    "name": 0,
    "email": 0,
    "phone": 0,
    "company": 0,
    "address": 0,
    "ssn": 0,
    "card": 0,
    "dob": 0,
    "ip": 0,
}


def get_redaction_summary():
    return {
        "name": int(replacement_counters.get("name", 0)),
        "email": int(replacement_counters.get("email", 0)),
        "phone": int(replacement_counters.get("phone", 0)),
        "company": int(replacement_counters.get("company", 0)),
        "address": int(replacement_counters.get("address", 0)),
        "ssn": int(replacement_counters.get("ssn", 0)),
        "card": int(replacement_counters.get("card", 0)),
        "dob": int(replacement_counters.get("dob", 0)),
        "ip": int(replacement_counters.get("ip", 0)),
    }


def reset_counters():
    for key in replacement_counters:
        replacement_counters[key] = 0


def get_replacement(category):

    replacement_counters[category] += 1

    return (
        f"[REDACTED_{category.upper()}_"
        f"{replacement_counters[category]}]"
    )


# ============================================================
# CREDIT CARD / LUHN
# ============================================================

def luhn_check(number):

    digits = [
        int(d)
        for d in number
        if d.isdigit()
    ]

    if not 13 <= len(digits) <= 19:
        return False

    total = 0
    parity = len(digits) % 2

    for i, digit in enumerate(digits):

        if i % 2 == parity:
            digit *= 2

            if digit > 9:
                digit -= 9

        total += digit

    return total % 10 == 0


# ============================================================
# ADDRESS HELPERS
# ============================================================

def looks_like_address(value):

    lower_value = value.lower()

    for word in ADDRESS_WORDS:

        if re.search(
            r"\b" + re.escape(word) + r"\b",
            lower_value,
        ):
            return True

    has_city = any(
        re.search(
            r"\b" + re.escape(city) + r"\b",
            lower_value,
        )
        for city in ADDRESS_CITIES
    )

    if has_city:
        return True

    has_state = any(
        re.search(
            r"\b" + re.escape(state) + r"\b",
            lower_value,
        )
        for state in INDIAN_STATES
    )

    if has_state:
        return True

    return False


def redact_addresses(text):

    """
    Detect addresses containing Indian PIN codes.

    Only a limited amount of text around each PIN is inspected.
    This prevents a PIN code from causing hundreds of words
    of prospectus text to be redacted.
    """

    matches = list(PIN_PATTERN.finditer(text))

    # Redact short number-led addresses without a PIN only when their address
    # wording and city make the classification sufficiently specific.
    for match in reversed(list(ADDRESS_WITHOUT_PIN_PATTERN.finditer(text))):
        if REDACTION_TOKEN_PATTERN.search(match.group(0)):
            continue

        text = (
            text[:match.start()]
            + get_replacement("address")
            + text[match.end():]
        )

    matches = list(PIN_PATTERN.finditer(text))

    if not matches:
        return text

    replacements = []

    for match in matches:

        start_pin = match.start()
        end_pin = match.end()

        # Look at limited text before PIN.
        start = max(0, start_pin - 160)

        # Look at limited text after PIN.
        end = min(len(text), end_pin + 80)

        candidate = text[start:end]

        # Do not touch existing redaction tokens.
        if REDACTION_TOKEN_PATTERN.search(candidate):
            continue

        if not looks_like_address(candidate):
            continue

        address_start = start

        separators = [
            "\n",
            "\t",
            "Telephone:",
            "Tel:",
            "Email:",
            "Website:",
            "Contact Person:",
        ]

        for separator in separators:

            pos = text.rfind(
                separator,
                start,
                start_pin,
            )

            if pos != -1:
                address_start = max(
                    address_start,
                    pos + len(separator),
                )

        address_start = max(
            address_start,
            start_pin - 140,
        )

        address_end = end_pin

        tail = text[
            end_pin:min(
                len(text),
                end_pin + 100,
            )
        ]

        state_match = re.search(
            r"^\s*,?\s*(?:"
            + "|".join(
                re.escape(x)
                for x in sorted(
                    INDIAN_STATES,
                    key=len,
                    reverse=True,
                )
            )
            + r")"
            r"(?:\s*,?\s*India)?",
            tail,
            re.IGNORECASE,
        )

        if state_match:
            address_end = (
                end_pin
                + state_match.end()
            )

        replacements.append(
            (
                address_start,
                address_end,
            )
        )

    # Merge overlapping ranges.
    merged = []

    for start, end in sorted(
        replacements
    ):

        if not merged:
            merged.append(
                [start, end]
            )
            continue

        previous = merged[-1]

        if start <= previous[1]:
            previous[1] = max(
                previous[1],
                end,
            )
        else:
            merged.append(
                [start, end]
            )

    # Replace right-to-left.
    for start, end in reversed(
        merged
    ):

        original = text[start:end]

        if REDACTION_TOKEN_PATTERN.fullmatch(
            original.strip()
        ):
            continue

        text = (
            text[:start]
            + get_replacement("address")
            + text[end:]
        )

    return text


# ============================================================
# REDACT TEXT
# ============================================================

def redact_text(text):

    if not text:
        return text

    # --------------------------------------------------------
    # Protect existing redaction tokens
    # --------------------------------------------------------

    protected_tokens = []

    def protect_token(match):

        index = len(
            protected_tokens
        )

        protected_tokens.append(
            match.group(0)
        )

        return (
            f"___REDACTION_TOKEN_{index}___"
        )

    text = REDACTION_TOKEN_PATTERN.sub(
        protect_token,
        text,
    )

    # --------------------------------------------------------
    # 1. EMAIL
    # --------------------------------------------------------

    text = EMAIL_PATTERN.sub(
        lambda m: get_replacement("email"),
        text,
    )

    # --------------------------------------------------------
    # 2. PHONE
    # --------------------------------------------------------

    text = PHONE_PATTERN.sub(
        lambda m: get_replacement("phone"),
        text,
    )

    # --------------------------------------------------------
    # 3. CREDIT CARD
    # --------------------------------------------------------

    def replace_card(match):

        value = match.group(0)

        if luhn_check(value):
            return get_replacement("card")

        return value

    text = CREDIT_CARD_PATTERN.sub(
        replace_card,
        text,
    )

    # --------------------------------------------------------
    # 4. SSN
    # --------------------------------------------------------

    text = SSN_PATTERN.sub(
        lambda m: get_replacement("ssn"),
        text,
    )

    # --------------------------------------------------------
    # 5. IP ADDRESS
    # --------------------------------------------------------

    text = IP_PATTERN.sub(
        lambda m: get_replacement("ip"),
        text,
    )

    # --------------------------------------------------------
    # 6. DATE OF BIRTH
    # --------------------------------------------------------

    def replace_dob(match):

        original_date = match.group(1)

        replacement = get_replacement(
            "dob"
        )

        return match.group(0).replace(
            original_date,
            replacement,
        )

    text = DOB_PATTERN.sub(
        replace_dob,
        text,
    )

    # --------------------------------------------------------
    # 7. PHYSICAL ADDRESS
    # --------------------------------------------------------

    text = redact_addresses(text)

    # --------------------------------------------------------
    # 8. KNOWN COMPANIES
    # --------------------------------------------------------

    for company in sorted(
        KNOWN_COMPANIES,
        key=len,
        reverse=True,
    ):

        pattern = re.compile(
            r"(?<![A-Za-z])"
            + re.escape(company)
            + r"(?![A-Za-z])",
            re.IGNORECASE,
        )

        text = pattern.sub(
            lambda m: get_replacement(
                "company"
            ),
            text,
        )

    # --------------------------------------------------------
    # 9. GENERIC COMPANY DETECTION
    # --------------------------------------------------------

    company_pattern = re.compile(
        r"""
        \b
        [A-Z][A-Za-z0-9&.,'()/-]*
        (?:
            \s+[A-Z][A-Za-z0-9&.,'()/-]*
        ){0,7}
        \s+
        (?:
            Private\s+Limited
            |
            Pvt\.?\s*Ltd\.?
            |
            Limited
            |
            LLP
            |
            Corporation
        )
        \b
        """,
        re.VERBOSE,
    )

    def replace_company(match):

        value = match.group(0).strip()

        if value.startswith(
            "[REDACTED_"
        ):
            return value

        return get_replacement(
            "company"
        )

    text = company_pattern.sub(
        replace_company,
        text,
    )

    # --------------------------------------------------------
    # 10. KNOWN PEOPLE
    # --------------------------------------------------------

    for name in sorted(
        KNOWN_NAMES,
        key=len,
        reverse=True,
    ):

        pattern = re.compile(
            r"(?<![A-Za-z])"
            + re.escape(name)
            + r"(?![A-Za-z])",
            re.IGNORECASE,
        )

        text = pattern.sub(
            lambda m: get_replacement(
                "name"
            ),
            text,
        )

    # --------------------------------------------------------
    # 11. SPACY PERSON DETECTION
    # --------------------------------------------------------

    ensure_nlp_ready()
    doc = nlp(text)

    entities = list(doc.ents)

    for entity in reversed(
        entities
    ):

        if entity.label_ != "PERSON":
            continue

        entity_text = entity.text.strip()

        if entity_text.startswith(
            "[REDACTED_"
        ):
            continue

        if len(
            entity_text.split()
        ) < 2:
            continue

        if any(
            char.isdigit()
            for char in entity_text
        ):
            continue

        if "@" in entity_text:
            continue

        if "http" in entity_text.lower():
            continue

        if "www." in entity_text.lower():
            continue

        lower_words = set(
            entity_text.lower().split()
        )

        # Avoid address/location entities.
        blocked_words = {
            "road",
            "street",
            "lane",
            "village",
            "nagar",
            "marg",
            "pune",
            "mumbai",
            "delhi",
            "building",
            "floor",
            "sector",
            "colony",
            "society",
            "house",
            "park",
            "tower",
            "complex",
            "office",
            "prabhadevi",
            "bandra",
            "vikhroli",
            "taluka",
            "pashan",
            "baner",
            "shivajinagar",
            "erandawane",
            "bhopal",
            "hyderabad",
            "chennai",
            "bangalore",
            "bengaluru",
        }

        if lower_words.intersection(
            blocked_words
        ):
            continue

        # Avoid company-like entities.
        company_words = {
            "limited",
            "private",
            "llp",
            "bank",
            "corporation",
            "industries",
            "company",
            "trust",
            "foundation",
        }

        if lower_words.intersection(
            company_words
        ):
            continue

        replacement = get_replacement(
            "name"
        )

        text = (
            text[:entity.start_char]
            + replacement
            + text[entity.end_char:]
        )

    # --------------------------------------------------------
    # Restore protected tokens
    # --------------------------------------------------------

    for index, original in enumerate(
        protected_tokens
    ):

        placeholder = (
            f"___REDACTION_TOKEN_{index}___"
        )

        text = text.replace(
            placeholder,
            original,
        )

    return text


# ============================================================
# PROCESS PARAGRAPH
# ============================================================

def process_paragraph(paragraph):

    original_text = paragraph.text

    if not original_text:
        return

    redacted_text = redact_text(
        original_text
    )

    if redacted_text != original_text:
        paragraph.text = redacted_text


# ============================================================
# PROCESS TABLE
# ============================================================

def process_table(table):

    for row in table.rows:

        for cell in row.cells:

            for paragraph in cell.paragraphs:
                process_paragraph(paragraph)

            for nested_table in cell.tables:
                process_table(nested_table)


# ============================================================
# PROCESS HEADERS AND FOOTERS
# ============================================================

def process_headers_and_footers(document):

    for section in document.sections:

        # Header
        header = section.header

        for paragraph in header.paragraphs:
            process_paragraph(paragraph)

        for table in header.tables:
            process_table(table)

        # Footer
        footer = section.footer

        for paragraph in footer.paragraphs:
            process_paragraph(paragraph)

        for table in footer.tables:
            process_table(table)


# ============================================================
# REDACT DOCUMENT
# ============================================================

def redact_document(
    input_file,
    output_file,
):

    ensure_nlp_ready()
    reset_counters()

    print("Reading document...")

    document = Document(
        input_file
    )

    print("Redacting paragraphs...")

    for paragraph in document.paragraphs:
        process_paragraph(paragraph)

    print("Redacting tables...")

    for table in document.tables:
        process_table(table)

    print("Redacting headers and footers...")

    process_headers_and_footers(
        document
    )

    output_file.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    print("Saving redacted document...")

    document.save(
        output_file
    )

    print()
    print("=" * 60)
    print("REDACTION COMPLETED")
    print("=" * 60)

    print(
        f"Input :  {input_file}"
    )

    print(
        f"Output:  {output_file}"
    )

    print()
    print("REDACTION SUMMARY")

    print(
        f"Name      : {replacement_counters['name']}"
    )

    print(
        f"Email     : {replacement_counters['email']}"
    )

    print(
        f"Phone     : {replacement_counters['phone']}"
    )

    print(
        f"Company   : {replacement_counters['company']}"
    )

    print(
        f"Address   : {replacement_counters['address']}"
    )

    print(
        f"SSN       : {replacement_counters['ssn']}"
    )

    print(
        f"Card      : {replacement_counters['card']}"
    )

    print(
        f"DOB       : {replacement_counters['dob']}"
    )

    print(
        f"IP        : {replacement_counters['ip']}"
    )

    print("=" * 60)


# ============================================================
# MAIN
# ============================================================

if __name__ == "__main__":

    if not INPUT_FILE.exists():

        print(
            "ERROR: Input document was not found."
        )

        print(INPUT_FILE)

        raise SystemExit(1)

    redact_document(
        INPUT_FILE,
        OUTPUT_FILE,
    )
