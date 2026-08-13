import re
from pathlib import Path
from docx import Document


BASE_DIR = Path(__file__).resolve().parent.parent
OUTPUT_FILE = BASE_DIR / "output" / "redacted_prospectus.docx"


EMAIL_PATTERN = re.compile(
    r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"
)

PHONE_PATTERN = re.compile(
    r"""
    (?<!\d)
    (?:
        \+?91[\s-]?
        (?:
            [6-9]\d{9}
            |
            \d{2}[\s-]?\d{4}[\s-]?\d{4}
            |
            \d{2}[\s-]?\d{8}
        )
        |
        [6-9]\d{9}
        |
        0\d{2,4}[\s-]\d{6,8}
    )
    (?!\d)
    """,
    re.VERBOSE
)

IP_PATTERN = re.compile(
    r"""
    (?<![\d.])
    (?:
        (?:25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)
        \.
        (?:25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)
        \.
        (?:25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)
        \.
        (?:25[0-5]|2[0-4]\d|1\d\d|[1-9]?\d)
    )
    (?![\d.])
    """,
    re.VERBOSE
)

SSN_PATTERN = re.compile(
    r"(?<!\d)(?!000|666|9\d\d)\d{3}[- ]\d{2}[- ]\d{4}(?!\d)"
)


def get_document_text(document):
    text_parts = []

    for paragraph in document.paragraphs:
        text_parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text_parts.append(paragraph.text)

    for section in document.sections:

        for paragraph in section.header.paragraphs:
            text_parts.append(paragraph.text)

        for paragraph in section.footer.paragraphs:
            text_parts.append(paragraph.text)

    return "\n".join(text_parts)


def verify_document():

    if not OUTPUT_FILE.exists():
        print("ERROR: Output document was not found.")
        print(OUTPUT_FILE)
        return

    print("Reading output document...")

    document = Document(OUTPUT_FILE)

    text = get_document_text(document)

    print()
    print("=" * 60)
    print("REDACTION VERIFICATION")
    print("=" * 60)

    emails = EMAIL_PATTERN.findall(text)
    phones = PHONE_PATTERN.findall(text)
    ips = IP_PATTERN.findall(text)
    ssns = SSN_PATTERN.findall(text)

    print()
    print("Emails found:", len(emails))

    for value in emails:
        print("  ", value)

    print()
    print("Phone numbers found:", len(phones))

    for value in phones:
        print("  ", value)

    print()
    print("IP addresses found:", len(ips))

    for value in ips:
        print("  ", value)

    print()
    print("SSNs found:", len(ssns))

    for value in ssns:
        print("  ", value)

    print()
    print("=" * 60)

    if not emails and not phones and not ips and not ssns:
        print("PASS: No detectable email, phone, IP, or SSN found.")
    else:
        print("WARNING: Potential PII still exists in the output.")

    print("=" * 60)


if __name__ == "__main__":
    verify_document()